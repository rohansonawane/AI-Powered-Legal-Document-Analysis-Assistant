import os
import magic
from pypdf import PdfReader
import docx
import streamlit as st
from typing import Dict, Any
import json
from datetime import datetime
import logging
from .image_processor import extract_images_from_pdf, extract_images_from_docx, get_image_description, process_image
import tempfile
import traceback
import cv2
import numpy as np
from pdf2image import convert_from_path

logger = logging.getLogger(__name__)

def process_document(file) -> Dict[str, Any]:
    """
    Process a document file and extract its content, metadata, and images.
    
    Args:
        file: The uploaded file object
        
    Returns:
        Dictionary containing document title, content, metadata, and images
    """
    try:
        # Get file extension
        file_ext = file.name.split('.')[-1].lower()
        
        # Create a temporary file to store the uploaded content
        with tempfile.NamedTemporaryFile(delete=False, suffix=f'.{file_ext}') as temp_file:
            temp_file.write(file.getvalue())
            temp_file_path = temp_file.name
        
        try:
            # Extract text and images based on file type
            if file_ext == 'pdf':
                result = process_pdf(temp_file_path)
            elif file_ext == 'docx':
                result = process_docx(temp_file_path)
            elif file_ext == 'txt':
                result = process_txt(temp_file_path)
            else:
                raise ValueError(f"Unsupported file type: {file_ext}")
            
            if not result or not result.get('content'):
                raise ValueError("No content extracted from document")
            
            # Clean and normalize the content
            content = clean_text(result['content'])
            if not content:
                raise ValueError("No valid text content after cleaning")
            
            # Extract metadata
            metadata = {
                'filename': file.name,
                'file_type': file_ext,
                'file_size': file.size,
                'processed_date': datetime.now().isoformat()
            }
            
            # Add image information to metadata if available
            if 'images' in result and result['images']:
                metadata['image_count'] = len(result['images'])
                try:
                    metadata['image_descriptions'] = [
                        get_image_description(img['image']) for img in result['images']
                    ]
                except Exception as img_error:
                    logger.warning(f"Error getting image descriptions: {str(img_error)}")
                    metadata['image_descriptions'] = []
            
            # Use the original filename as the title
            title = file.name
            
            logger.info(f"Processed document: {title}")
            logger.info(f"Content length: {len(content)} characters")
            if 'images' in result:
                logger.info(f"Found {len(result['images'])} images")
            
            return {
                'title': title,
                'content': content,
                'metadata': metadata,
                'images': result.get('images', [])
            }
            
        finally:
            # Clean up the temporary file
            try:
                os.unlink(temp_file_path)
            except Exception as e:
                logger.warning(f"Failed to delete temporary file {temp_file_path}: {str(e)}")
        
    except Exception as e:
        logger.error(f"Error processing document: {str(e)}")
        logger.error(f"Error details: {traceback.format_exc()}")
        st.error(f"Error processing document: {str(e)}")
        return None

def process_pdf(file_path) -> dict:
    """Process PDF file and extract text content and images."""
    try:
        # Read PDF
        pdf = PdfReader(file_path)
        
        # Extract text from each page with progress tracking
        text_content = []
        total_pages = len(pdf.pages)
        
        if total_pages == 0:
            raise ValueError("PDF file is empty or corrupted")
        
        # Process pages in batches of 5 for better memory management
        batch_size = 5
        for i in range(0, total_pages, batch_size):
            batch_end = min(i + batch_size, total_pages)
            st.write(f"Processing pages {i+1} to {batch_end} of {total_pages}...")
            
            for page_num in range(i, batch_end):
                try:
                    page = pdf.pages[page_num]
                    
                    # Try direct text extraction first
                    text = page.extract_text()
                    
                    # If no text found, try OCR
                    if not text or not text.strip():
                        logger.info(f"No text found on page {page_num + 1}, trying OCR...")
                        try:
                            # Convert page to image
                            images = convert_from_path(file_path, first_page=page_num+1, last_page=page_num+1)
                            if images:
                                # Convert PIL image to OpenCV format
                                opencv_image = cv2.cvtColor(np.array(images[0]), cv2.COLOR_RGB2BGR)
                                # Process image and extract text
                                _, text = process_image(opencv_image)
                                logger.info(f"Successfully extracted text using OCR from page {page_num + 1}")
                        except Exception as ocr_error:
                            logger.warning(f"OCR failed for page {page_num + 1}: {str(ocr_error)}")
                            continue
                    
                    if text and text.strip():  # Only add non-empty pages
                        text_content.append(text)
                        logger.info(f"Successfully extracted text from page {page_num + 1}")
                    else:
                        logger.warning(f"No text could be extracted from page {page_num + 1}")
                        
                except Exception as page_error:
                    logger.warning(f"Error processing page {page_num + 1}: {str(page_error)}")
                    continue
        
        if not text_content:
            raise ValueError("No text content could be extracted from PDF")
        
        # Combine all text
        full_text = '\n\n'.join(text_content)
        
        # Clean up text
        cleaned_text = clean_text(full_text)
        
        if not cleaned_text:
            raise ValueError("No valid text content after cleaning")
        
        # Extract images
        try:
            images = extract_images_from_pdf(file_path)
        except Exception as img_error:
            logger.warning(f"Error extracting images: {str(img_error)}")
            images = []
        
        return {
            'title': os.path.splitext(os.path.basename(file_path))[0],
            'content': cleaned_text,
            'metadata': {
                'file_type': 'pdf',
                'page_count': total_pages,
                'file_name': os.path.basename(file_path),
                'processed_pages': len(text_content)
            },
            'images': images
        }
    except Exception as e:
        logger.error(f"Error processing PDF: {str(e)}")
        logger.error(f"Error details: {traceback.format_exc()}")
        st.error(f"Error processing PDF: {str(e)}")
        return None

def process_docx(file_path) -> dict:
    """Process DOCX file and extract text content and images."""
    try:
        doc = docx.Document(file_path)
        
        # Extract text from paragraphs
        text_content = []
        for para in doc.paragraphs:
            if para.text.strip():  # Only add non-empty paragraphs
                text_content.append(para.text)
        
        if not text_content:
            raise ValueError("No text content could be extracted from DOCX")
        
        # Combine all text
        full_text = '\n\n'.join(text_content)
        
        # Clean up text
        cleaned_text = clean_text(full_text)
        
        if not cleaned_text:
            raise ValueError("No valid text content after cleaning")
        
        # Extract images
        try:
            images = extract_images_from_docx(file_path)
        except Exception as img_error:
            logger.warning(f"Error extracting images: {str(img_error)}")
            images = []
        
        return {
            'title': os.path.splitext(os.path.basename(file_path))[0],
            'content': cleaned_text,
            'metadata': {
                'file_type': 'docx',
                'paragraph_count': len(doc.paragraphs),
                'file_name': os.path.basename(file_path)
            },
            'images': images
        }
    except Exception as e:
        logger.error(f"Error processing DOCX: {str(e)}")
        logger.error(f"Error details: {traceback.format_exc()}")
        st.error(f"Error processing DOCX: {str(e)}")
        return None

def process_txt(file_path) -> dict:
    """Process TXT file and extract text content."""
    try:
        # Read text content
        with open(file_path, 'r', encoding='utf-8') as f:
            text_content = f.read()
        
        if not text_content:
            raise ValueError("TXT file is empty")
        
        # Clean up text
        cleaned_text = clean_text(text_content)
        
        if not cleaned_text:
            raise ValueError("No valid text content after cleaning")
        
        return {
            'title': os.path.splitext(os.path.basename(file_path))[0],
            'content': cleaned_text,
            'metadata': {
                'file_type': 'txt',
                'file_name': os.path.basename(file_path)
            }
        }
    except Exception as e:
        logger.error(f"Error processing TXT: {str(e)}")
        logger.error(f"Error details: {traceback.format_exc()}")
        st.error(f"Error processing TXT: {str(e)}")
        return None

def clean_text(text: str) -> str:
    """
    Clean and normalize text content.
    
    Args:
        text: Raw text content
        
    Returns:
        Cleaned text
    """
    if not text:
        return ""
        
    # Replace multiple newlines with single newline
    text = '\n'.join(line.strip() for line in text.split('\n') if line.strip())
    
    # Replace multiple spaces with single space
    text = ' '.join(text.split())
    
    # Remove special characters but keep basic punctuation
    text = ''.join(char for char in text if char.isprintable() or char in '.,;:!?()[]{}')
    
    return text 